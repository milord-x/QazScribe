from datetime import datetime
from html import escape
import json
from pathlib import Path
import textwrap
from typing import Any, Callable

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


class DocumentGenerationError(RuntimeError):
    """Raised when result documents cannot be generated."""


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return "\n".join(f"- {item}" for item in value)
    return str(value)


def _section(title: str, body: Any) -> str:
    content = _text(body).strip() or "Not available."
    return f"\n{title}\n{'=' * len(title)}\n{content}\n"


DOCUMENT_FORMATS = {"txt", "srt", "vtt", "json", "html", "docx", "pdf"}
DOCUMENT_FILENAMES = {
    "txt": "qtranscript_result.txt",
    "srt": "qtranscript_result.srt",
    "vtt": "qtranscript_result.vtt",
    "json": "qtranscript_result.json",
    "html": "qtranscript_result.html",
    "docx": "qtranscript_result.docx",
    "pdf": "qtranscript_result.pdf",
}
LANGUAGE_LABELS = {
    "ru": {
        "kk": "Казахский",
        "ky": "Кыргызский",
    },
    "kk": {
        "kk": "Қазақша",
        "ky": "Қырғызша",
    },
    "en": {
        "kk": "Kazakh",
        "ky": "Kyrgyz",
    },
}
DOCUMENT_LABELS = {
    "ru": {
        "title": "Qtranscript: текст записи",
        "meta": "Сведения",
        "task_id": "Номер задачи",
        "recording_started_at": "Время начала записи",
        "recording_duration": "Длительность записи",
        "speech_language": "Язык речи",
        "original_text": "Оригинальная речь",
        "translated_text": "Переведённый текст",
        "not_available": "Нет данных.",
        "unknown_language": "не определён",
    },
    "kk": {
        "title": "Qtranscript: жазба мәтіні",
        "meta": "Мәлімет",
        "task_id": "Тапсырма нөмірі",
        "recording_started_at": "Жазба басталған уақыт",
        "recording_duration": "Жазба ұзақтығы",
        "speech_language": "Сөйлеу тілі",
        "original_text": "Түпнұсқа сөз",
        "translated_text": "Аударылған мәтін",
        "not_available": "Дерек жоқ.",
        "unknown_language": "анықталмаған",
    },
    "en": {
        "title": "Qtranscript: recording text",
        "meta": "Details",
        "task_id": "Task number",
        "recording_started_at": "Recording start",
        "recording_duration": "Recording duration",
        "speech_language": "Speech language",
        "original_text": "Original speech",
        "translated_text": "Translated text",
        "not_available": "Not available.",
        "unknown_language": "unknown",
    },
}


def _short_task_id(task_id: str) -> str:
    return task_id.split("-", 1)[0] or task_id


def _ui_language(language_code: str | None) -> str:
    normalized = (language_code or "").strip().lower()
    return normalized if normalized in DOCUMENT_LABELS else "ru"


def _labels(language_code: str | None) -> dict[str, str]:
    return DOCUMENT_LABELS[_ui_language(language_code)]


def _language_label(language_code: str | None, ui_language: str | None = None) -> str:
    normalized = (language_code or "").strip().lower()
    language_labels = LANGUAGE_LABELS[_ui_language(ui_language)]
    return language_labels.get(normalized, normalized or _labels(ui_language)["unknown_language"])


def _format_recording_started(value: str | None) -> str:
    raw_value = (value or "").strip()
    if not raw_value:
        return "unknown"
    try:
        normalized = raw_value.replace("Z", "+00:00")
        return datetime.fromisoformat(normalized).strftime("%Y-%m-%d %H:%M")
    except ValueError:
        return raw_value


def _font_name() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/liberation/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            font_name = path.stem
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
    return "Helvetica"


def _build_payload(
    task_id: str,
    detected_language: str | None,
    transcript: dict[str, Any],
    translation: dict[str, Any],
    summary: dict[str, Any],
    metadata: dict[str, Any] | None = None,
) -> dict[str, Any]:
    metadata = metadata or {}
    ui_language = _ui_language(str(metadata.get("ui_language") or "ru"))
    labels = _labels(ui_language)
    segments = _subtitle_segments(transcript)
    speaker_names = sorted({segment.get("speaker", "Спикер 1") for segment in segments})
    duration = metadata.get("recording_duration_seconds") or _recording_duration(segments)
    target_language = translation.get("target_language") or "kk"
    detected_language_value = detected_language or transcript.get("detected_language")
    return {
        "title": labels["title"],
        "labels": labels,
        "task_id": task_id,
        "display_task_id": metadata.get("display_task_id") or _short_task_id(task_id),
        "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "recording_started_at": _format_recording_started(metadata.get("recording_started_at")),
        "recording_duration": _format_duration(duration),
        "speaker_count": len(speaker_names) if speaker_names else 0,
        "speaker_names": speaker_names,
        "detected_language": detected_language_value or "unknown",
        "detected_language_label": _language_label(detected_language_value, ui_language),
        "translation_target_language": target_language,
        "translation_title": labels["translated_text"],
        "original_transcript": transcript.get("full_transcript") or "",
        "speaker_transcript": _speaker_transcript(segments),
        "translated_text": translation.get("translated_text") or "",
        "short_summary": summary.get("short_summary") or "",
        "detailed_summary": summary.get("detailed_summary") or "",
        "key_points": summary.get("key_points") or [],
        "decisions": summary.get("decisions") or [],
        "action_items": summary.get("action_items") or [],
        "participants_or_speakers": summary.get("participants_or_speakers") or [],
        "risks_or_open_questions": summary.get("risks_or_open_questions") or [],
        "final_notes": summary.get("final_notes") or "",
    }


def _plain_text(payload: dict[str, Any]) -> str:
    labels = payload["labels"]
    parts = [
        payload["title"],
        _section(
            labels["meta"],
            "\n".join(
                [
                    f"{labels['task_id']}: {payload['display_task_id']}",
                    f"{labels['recording_started_at']}: {payload['recording_started_at']}",
                    f"{labels['recording_duration']}: {payload['recording_duration']}",
                    f"{labels['speech_language']}: {payload['detected_language_label']}",
                ]
            ),
        ),
        _section(labels["original_text"], payload["original_transcript"]),
        _section(payload["translation_title"], payload["translated_text"]),
    ]
    return "\n".join(parts).strip() + "\n"


def _html(payload: dict[str, Any]) -> str:
    labels = payload["labels"]

    def list_or_paragraph(value: Any) -> str:
        if isinstance(value, list):
            if not value:
                return f"<p>{escape(labels['not_available'])}</p>"
            return "<ul>" + "".join(f"<li>{escape(str(item))}</li>" for item in value) + "</ul>"
        return f"<p>{escape(_text(value).strip() or labels['not_available'])}</p>"

    sections = [
        (labels["original_text"], payload["original_transcript"]),
        (payload["translation_title"], payload["translated_text"]),
    ]
    body = "\n".join(
        f"<section><h2>{escape(title)}</h2>{list_or_paragraph(value)}</section>"
        for title, value in sections
    )
    return f"""<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8" />
    <title>{escape(payload["title"])}</title>
    <style>
      body {{ font-family: Arial, sans-serif; line-height: 1.5; margin: 32px; color: #18202a; }}
      h1 {{ margin-bottom: 8px; }}
      h2 {{ margin-top: 28px; border-bottom: 1px solid #d9dee7; padding-bottom: 6px; }}
      p, li {{ white-space: pre-wrap; }}
      .meta {{ color: #667085; }}
    </style>
  </head>
  <body>
    <h1>{escape(payload["title"])}</h1>
    <p class="meta">{escape(labels["task_id"])}: {escape(payload["display_task_id"])}</p>
    <p class="meta">{escape(labels["recording_started_at"])}: {escape(payload["recording_started_at"])}</p>
    <p class="meta">{escape(labels["recording_duration"])}: {escape(payload["recording_duration"])}</p>
    <p class="meta">{escape(labels["speech_language"])}: {escape(payload["detected_language_label"])}</p>
    {body}
  </body>
</html>
"""


def _timestamp_srt(seconds: float) -> str:
    milliseconds = int(round(seconds * 1000))
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def _timestamp_vtt(seconds: float) -> str:
    return _timestamp_srt(seconds).replace(",", ".")


def _subtitle_segments(transcript: dict[str, Any]) -> list[dict[str, Any]]:
    segments = transcript.get("segments") or []
    normalized = []
    for index, segment in enumerate(segments, start=1):
        text = str(segment.get("text") or "").strip()
        if not text:
            continue
        start = float(segment.get("start") or 0)
        end = float(segment.get("end") or max(start + 1, index))
        normalized.append(
            {
                "start": start,
                "end": max(end, start + 0.5),
                "text": text,
                "speaker": str(segment.get("speaker") or "Спикер 1"),
            }
        )
    if normalized:
        return normalized

    full_text = str(transcript.get("full_transcript") or "").strip()
    return [{"start": 0.0, "end": 3.0, "text": full_text, "speaker": "Спикер 1"}] if full_text else []


def _recording_duration(segments: list[dict[str, Any]]) -> float:
    if not segments:
        return 0.0
    return max(float(segment.get("end") or 0) for segment in segments)


def _format_duration(seconds: float | int | None) -> str:
    total_seconds = int(round(float(seconds or 0)))
    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours:02}:{minutes:02}:{secs:02}"
    return f"{minutes:02}:{secs:02}"


def _speaker_transcript(segments: list[dict[str, Any]]) -> str:
    lines = []
    for segment in segments:
        lines.append(
            f"[{_timestamp_vtt(float(segment['start']))}] {segment.get('speaker', 'Спикер 1')}: {segment['text']}"
        )
    return "\n".join(lines)


def _srt(transcript: dict[str, Any]) -> str:
    blocks = []
    for index, segment in enumerate(_subtitle_segments(transcript), start=1):
        blocks.append(
            "\n".join(
                [
                    str(index),
                    f"{_timestamp_srt(segment['start'])} --> {_timestamp_srt(segment['end'])}",
                    f"{segment.get('speaker', 'Спикер 1')}: {segment['text']}",
                ]
            )
        )
    return "\n\n".join(blocks).strip() + "\n"


def _vtt(transcript: dict[str, Any]) -> str:
    blocks = ["WEBVTT"]
    for segment in _subtitle_segments(transcript):
        blocks.append(
            "\n".join(
                [
                    f"{_timestamp_vtt(segment['start'])} --> {_timestamp_vtt(segment['end'])}",
                    f"{segment.get('speaker', 'Спикер 1')}: {segment['text']}",
                ]
            )
        )
    return "\n\n".join(blocks).strip() + "\n"


def _write_docx(path: Path, payload: dict[str, Any]) -> None:
    labels = payload["labels"]
    document = Document()
    document.add_heading(payload["title"], level=0)
    document.add_paragraph(f"{labels['task_id']}: {payload['display_task_id']}")
    document.add_paragraph(f"{labels['recording_started_at']}: {payload['recording_started_at']}")
    document.add_paragraph(f"{labels['recording_duration']}: {payload['recording_duration']}")
    document.add_paragraph(f"{labels['speech_language']}: {payload['detected_language_label']}")

    for title, key in [
        (labels["original_text"], "original_transcript"),
        (payload["translation_title"], "translated_text"),
    ]:
        document.add_heading(title, level=1)
        value = payload[key]
        if isinstance(value, list):
            if not value:
                document.add_paragraph("Not available.")
            for item in value:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(_text(value).strip() or labels["not_available"])

    document.save(path)


def _write_pdf(path: Path, payload: dict[str, Any]) -> None:
    labels = payload["labels"]
    font_name = _font_name()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "QtranscriptNormal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    heading = ParagraphStyle(
        "QtranscriptHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceBefore=12,
        spaceAfter=6,
    )
    title = ParagraphStyle(
        "QtranscriptTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=18,
        leading=22,
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story = [
        Paragraph(escape(payload["title"]), title),
        Paragraph(escape(f"{labels['task_id']}: {payload['display_task_id']}"), normal),
        Paragraph(escape(f"{labels['recording_started_at']}: {payload['recording_started_at']}"), normal),
        Paragraph(escape(f"{labels['recording_duration']}: {payload['recording_duration']}"), normal),
        Paragraph(escape(f"{labels['speech_language']}: {payload['detected_language_label']}"), normal),
        Spacer(1, 8),
    ]

    for section_title, key in [
        (labels["original_text"], "original_transcript"),
        (payload["translation_title"], "translated_text"),
    ]:
        story.append(Paragraph(escape(section_title), heading))
        text = _text(payload[key]).strip() or labels["not_available"]
        for chunk in textwrap.wrap(text, width=180, replace_whitespace=False) or [text]:
            story.append(Paragraph(escape(chunk), normal))

    doc.build(story)


def generate_documents(
    task_id: str,
    output_dir: Path,
    detected_language: str | None,
    transcript: dict[str, Any],
    translation: dict[str, Any],
    summary: dict[str, Any],
    metadata: dict[str, Any] | None = None,
    ui_language: str | None = None,
) -> dict[str, Path]:
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        paths = {
            file_format: output_dir / filename
            for file_format, filename in DOCUMENT_FILENAMES.items()
        }

        for file_format, path in paths.items():
            generate_document(
                task_id,
                path,
                file_format,
                detected_language,
                transcript,
                translation,
                summary,
                _metadata_with_ui_language(metadata, ui_language),
            )
        return paths
    except Exception as exc:
        raise DocumentGenerationError(f"Document generation failed: {exc}") from exc


def generate_document(
    task_id: str,
    path: Path,
    file_format: str,
    detected_language: str | None,
    transcript: dict[str, Any],
    translation: dict[str, Any],
    summary: dict[str, Any],
    metadata: dict[str, Any] | None = None,
    ui_language: str | None = None,
) -> Path:
    if file_format not in DOCUMENT_FORMATS:
        raise DocumentGenerationError(f"Unsupported document format: {file_format}")

    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        payload = _build_payload(
            task_id,
            detected_language,
            transcript,
            translation,
            summary,
            _metadata_with_ui_language(metadata, ui_language),
        )
        writers: dict[str, Callable[[], None]] = {
            "txt": lambda: path.write_text(_plain_text(payload), encoding="utf-8"),
            "srt": lambda: path.write_text(_srt(transcript), encoding="utf-8"),
            "vtt": lambda: path.write_text(_vtt(transcript), encoding="utf-8"),
            "json": lambda: path.write_text(
                json.dumps(
                    {
                        "task_id": task_id,
                        "detected_language": detected_language,
                        "transcript": transcript,
                        "translation": translation,
                        "summary": summary,
                        "metadata": metadata or {},
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            ),
            "html": lambda: path.write_text(_html(payload), encoding="utf-8"),
            "docx": lambda: _write_docx(path, payload),
            "pdf": lambda: _write_pdf(path, payload),
        }
        writers[file_format]()
        return path
    except Exception as exc:
        raise DocumentGenerationError(f"Document generation failed: {exc}") from exc


def _metadata_with_ui_language(
    metadata: dict[str, Any] | None,
    ui_language: str | None,
) -> dict[str, Any]:
    prepared = dict(metadata or {})
    if ui_language:
        prepared["ui_language"] = ui_language
    return prepared
